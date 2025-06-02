from flask import Flask, request, send_file, render_template_string, session
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image
import io, zipfile, os, uuid

app = Flask(__name__)
app.secret_key = 'your-secret-key'

TEMP_FOLDER = "/tmp"
os.makedirs(TEMP_FOLDER, exist_ok=True)

HTML_PAGE = '''
<!DOCTYPE html>
<html>
<head>
    <title>Folder to Word</title>
    <style>
        body {
            font-family: 'Segoe UI', sans-serif;
            background: #f2f2f2;
            margin: 0;
            padding: 0;
        }
        .container {
            max-width: 600px;
            margin: 5em auto;
            padding: 2em;
            background: white;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
            border-radius: 12px;
            text-align: center;
        }
        h1 {
            font-size: 2em;
            margin-bottom: 1em;
        }
        input[type="file"] {
            margin: 1.5em 0;
            font-size: 1.1em;
        }
        button {
            padding: 15px 30px;
            font-size: 1.2em;
            background-color: #007BFF;
            color: white;
            border: none;
            border-radius: 8px;
            cursor: pointer;
        }
        button:hover {
            background-color: #0056b3;
        }
        .message {
            margin-top: 2em;
            color: green;
            font-size: 1.2em;
        }
        a.download-button {
            display: inline-block;
            margin-top: 1.5em;
            padding: 14px 28px;
            font-size: 1.2em;
            background-color: #28a745;
            color: white;
            border-radius: 8px;
            text-decoration: none;
        }
        a.download-button:hover {
            background-color: #1e7e34;
        }
        #doneMsg {
            font-size: 1.2em;
            margin-top: 2em;
            color: #333;
            display: none;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>🖼️ Folder to Word Document</h1>
        <form method="POST" enctype="multipart/form-data">
            <input type="file" name="zipfile" accept=".zip" required>
            <br>
            <button type="submit">Upload & Generate</button>
        </form>
        {% if file_ready %}
            <div class="message">
                ✅ Document ready!
                <br>
                <a href="/download/{{ doc_id }}" onclick="afterDownload()" class="download-button" download>
                    📥 Download Word Document
                </a>
            </div>
        {% endif %}
        <div class="message" id="doneMsg">🎉 Downloaded! Reloading...</div>
    </div>
    <script>
        function afterDownload() {
            document.getElementById('doneMsg').style.display = 'block';
            setTimeout(() => { window.location.href = "/" }, 3000);
        }
    </script>
</body>
</html>
'''

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        uploaded_file = request.files.get("zipfile")
        if not uploaded_file or not uploaded_file.filename.endswith(".zip"):
            return "Invalid ZIP file", 400

        zip_bytes = io.BytesIO(uploaded_file.read())

        try:
            with zipfile.ZipFile(zip_bytes) as zip_ref:
                file_list = zip_ref.namelist()
                extracted = {
                    name: zip_ref.read(name) for name in file_list
                    if not name.endswith("/") and name.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif'))
                }

                top_folders = [f.split('/')[0] for f in file_list if '/' in f]
                main_folder_name = top_folders[0] if top_folders else "images"

            folders = {}
            for path, data in extracted.items():
                folder = os.path.dirname(path)
                folders.setdefault(folder, []).append((path, data))

            doc = Document()
            doc.add_heading("📷 Image Gallery", 0)

            for folder, files in folders.items():
                heading = doc.add_paragraph()
                run = heading.add_run(folder or "Root")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

                for filename, img_bytes in sorted(files):
                    try:
                        caption = doc.add_paragraph()
                        run = caption.add_run(os.path.basename(filename))
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

                        image_stream = io.BytesIO(img_bytes)
                        Image.open(image_stream).verify()
                        image_stream.seek(0)
                        doc.add_picture(image_stream, width=Inches(5.5))
                        doc.add_paragraph("\n\n\n")
                    except Exception as e:
                        doc.add_paragraph(f"[Error inserting {filename}: {e}]")

            doc_id = str(uuid.uuid4())
            file_path = os.path.join(TEMP_FOLDER, f"{doc_id}.docx")
            doc.save(file_path)

            session['doc_id'] = doc_id
            session['file_name'] = f"{main_folder_name}.docx"

            return render_template_string(HTML_PAGE, file_ready=True, doc_id=doc_id)

        except Exception as e:
            return f"Error processing ZIP: {e}", 500

    return render_template_string(HTML_PAGE, file_ready=False)

@app.route("/download/<doc_id>")
def download(doc_id):
    file_path = os.path.join(TEMP_FOLDER, f"{doc_id}.docx")
    if not os.path.exists(file_path):
        return "Document expired or not found", 404

    file_name = session.get('file_name', 'output.docx')
    return send_file(file_path, as_attachment=True, download_name=file_name)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    print(f"Starting app on port {port}")
    app.run(host="0.0.0.0", port=port)
